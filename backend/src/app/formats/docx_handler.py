from pathlib import Path

from docx import Document

from app.formats.base import FormatHandler, register
from app.models.job import TextSegment


@register([".docx"])
class DocxHandler(FormatHandler):
    def extract_texts(self, file_path: Path) -> list[TextSegment]:
        doc = Document(str(file_path))
        segments = []
        for p_idx, paragraph in enumerate(doc.paragraphs):
            for r_idx, run in enumerate(paragraph.runs):
                if run.text.strip():
                    segments.append(TextSegment(
                        id=f"p{p_idx}:r{r_idx}",
                        text=run.text,
                        context=paragraph.text[:200],
                        metadata={"paragraph": p_idx, "run": r_idx},
                    ))
        return segments

    def replace_texts(self, file_path: Path, translated: list[TextSegment], output_path: Path) -> None:
        doc = Document(str(file_path))
        lookup = {seg.id: seg.text for seg in translated}

        for p_idx, paragraph in enumerate(doc.paragraphs):
            for r_idx, run in enumerate(paragraph.runs):
                key = f"p{p_idx}:r{r_idx}"
                if key in lookup:
                    run.text = lookup[key]

        doc.save(str(output_path))
